#!/usr/bin/env python3
"""Generate a professional cover letter in DOCX format.

Usage:
    python generate_cover_letter.py --input cover_data.json --output cover_letter.docx

Input JSON schema:
{
    "applicant": {
        "name": "John Doe",
        "email": "john@example.com",
        "phone": "+1 (555) 123-4567",
        "location": "San Francisco, CA",
        "linkedin": "linkedin.com/in/johndoe"
    },
    "recipient": {
        "name": "Jane Smith",
        "title": "Engineering Manager",
        "company": "InnovateTech",
        "address": "123 Market St, San Francisco, CA 94105"
    },
    "position": "Senior Software Engineer",
    "date": "February 25, 2026",
    "opening": "I am writing to express my strong interest in the Senior Software Engineer position...",
    "body_paragraphs": [
        "In my current role at TechCorp, I led the migration of three legacy systems...",
        "What particularly excites me about InnovateTech is your commitment to..."
    ],
    "closing": "I would welcome the opportunity to discuss how my experience...",
    "referral_source": "LinkedIn posting"
}
"""

import argparse
import json
import sys
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx required. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def generate_cover_letter(data, output_path):
    doc = Document()

    # Set up styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    applicant = data.get('applicant', {})
    recipient = data.get('recipient', {})

    # Applicant header
    header_p = doc.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    name_run = header_p.add_run(applicant.get('name', ''))
    name_run.bold = True
    name_run.font.size = Pt(14)
    name_run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    contact_parts = []
    for field in ['email', 'phone', 'location', 'linkedin']:
        if applicant.get(field):
            contact_parts.append(applicant[field])

    if contact_parts:
        contact_p = doc.add_paragraph(' | '.join(contact_parts))
        contact_p.runs[0].font.size = Pt(9.5)
        contact_p.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Date
    date_str = data.get('date', datetime.now().strftime('%B %d, %Y'))
    doc.add_paragraph('')  # spacer
    doc.add_paragraph(date_str)

    # Recipient
    if recipient:
        if recipient.get('name'):
            doc.add_paragraph(recipient['name'])
        if recipient.get('title'):
            doc.add_paragraph(recipient['title'])
        if recipient.get('company'):
            doc.add_paragraph(recipient['company'])
        if recipient.get('address'):
            doc.add_paragraph(recipient['address'])

    doc.add_paragraph('')  # spacer

    # Salutation
    greeting_name = recipient.get('name', 'Hiring Manager')
    doc.add_paragraph(f"Dear {greeting_name},")

    doc.add_paragraph('')

    # Opening paragraph
    if data.get('opening'):
        doc.add_paragraph(data['opening'])

    # Body paragraphs
    for para in data.get('body_paragraphs', []):
        doc.add_paragraph(para)

    # Closing paragraph
    if data.get('closing'):
        doc.add_paragraph(data['closing'])

    doc.add_paragraph('')

    # Sign-off
    doc.add_paragraph('Sincerely,')
    doc.add_paragraph('')
    sign_p = doc.add_paragraph(applicant.get('name', ''))
    sign_p.runs[0].bold = True

    doc.save(output_path)

    # Count words for validation
    word_count = sum(
        len(p.text.split()) for p in doc.paragraphs if p.text.strip()
    )

    return {
        "status": "success",
        "output": output_path,
        "word_count": word_count,
        "warning": "Cover letter exceeds recommended 400 words" if word_count > 400 else None
    }


def main():
    parser = argparse.ArgumentParser(description='Generate professional cover letter')
    parser.add_argument('--input', required=True, help='JSON file with cover letter data')
    parser.add_argument('--output', required=True, help='Output DOCX file path')
    args = parser.parse_args()

    with open(args.input, 'r') as f:
        data = json.load(f)

    result = generate_cover_letter(data, args.output)
    print(json.dumps(result))


if __name__ == '__main__':
    main()
