#!/usr/bin/env python3
"""Generate a professional biography in DOCX format.

Usage:
    python generate_bio.py --input bio_data.json --output bio.docx [--length short|medium|long]

Input JSON schema:
{
    "name": "John Doe",
    "current_title": "Senior Software Engineer",
    "current_company": "TechCorp",
    "years_experience": 8,
    "specializations": ["cloud architecture", "distributed systems", "team leadership"],
    "key_achievements": [
        "Led cloud migration saving $2M annually",
        "Built platform serving 10M daily users",
        "Open-source contributor with 5K+ GitHub stars"
    ],
    "education": "B.S. Computer Science, UC Berkeley",
    "certifications": ["AWS Solutions Architect", "Google Cloud Professional"],
    "speaking": ["KubeCon 2024", "PyCon 2023"],
    "personal_note": "passionate about mentoring and open source",
    "contact": {
        "email": "john@example.com",
        "linkedin": "linkedin.com/in/johndoe",
        "website": "johndoe.com"
    }
}

Output lengths:
- short: 50-75 words (LinkedIn headline, speaker bio)
- medium: 100-150 words (conference bio, website about)
- long: 200-300 words (full professional bio)
"""

import argparse
import json
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx required. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def build_short_bio(data):
    """50-75 word bio for speaker introductions, social media."""
    parts = []
    parts.append(f"{data['name']} is a {data.get('current_title', 'professional')}")
    if data.get('current_company'):
        parts.append(f"at {data['current_company']}")
    if data.get('specializations'):
        specs = data['specializations'][:2]
        parts.append(f"specializing in {' and '.join(specs)}")
    parts.append('.')
    if data.get('key_achievements'):
        parts.append(data['key_achievements'][0] + '.')
    if data.get('personal_note'):
        parts.append(f"{data['name'].split()[0]} is {data['personal_note']}.")
    return ' '.join(parts)


def build_medium_bio(data):
    """100-150 word bio for conferences, websites."""
    paragraphs = []

    intro = f"{data['name']} is a {data.get('current_title', 'professional')}"
    if data.get('current_company'):
        intro += f" at {data['current_company']}"
    if data.get('years_experience'):
        intro += f" with {data['years_experience']}+ years of experience"
    if data.get('specializations'):
        intro += f" in {', '.join(data['specializations'][:3])}"
    intro += '.'
    paragraphs.append(intro)

    if data.get('key_achievements'):
        achievements = '. '.join(data['key_achievements'][:2])
        paragraphs.append(achievements + '.')

    if data.get('education'):
        paragraphs.append(f"{data['name'].split()[0]} holds a {data['education']}.")

    if data.get('personal_note'):
        paragraphs.append(f"Outside of work, {data['name'].split()[0].lower()} is {data['personal_note']}.")

    return '\n\n'.join(paragraphs)


def build_long_bio(data):
    """200-300 word full professional bio."""
    paragraphs = []

    intro = f"{data['name']} is a {data.get('current_title', 'professional')}"
    if data.get('current_company'):
        intro += f" at {data['current_company']}"
    if data.get('years_experience'):
        intro += f" with over {data['years_experience']} years of experience"
    if data.get('specializations'):
        intro += f" specializing in {', '.join(data['specializations'])}"
    intro += '.'
    paragraphs.append(intro)

    if data.get('key_achievements'):
        achievement_text = "Notable accomplishments include: " + '; '.join(data['key_achievements']) + '.'
        paragraphs.append(achievement_text)

    if data.get('education') or data.get('certifications'):
        edu_parts = []
        if data.get('education'):
            edu_parts.append(f"{data['name'].split()[0]} holds a {data['education']}")
        if data.get('certifications'):
            edu_parts.append(f"is certified as a {', '.join(data['certifications'])}")
        paragraphs.append(' and '.join(edu_parts) + '.')

    if data.get('speaking'):
        paragraphs.append(
            f"A recognized voice in the industry, {data['name'].split()[0]} has spoken at "
            f"{', '.join(data['speaking'])}."
        )

    if data.get('personal_note'):
        paragraphs.append(
            f"Beyond professional work, {data['name'].split()[0]} is {data['personal_note']}."
        )

    if data.get('contact'):
        contact = data['contact']
        contact_parts = []
        if contact.get('linkedin'):
            contact_parts.append(contact['linkedin'])
        if contact.get('website'):
            contact_parts.append(contact['website'])
        if contact_parts:
            paragraphs.append(f"Connect at {' or '.join(contact_parts)}.")

    return '\n\n'.join(paragraphs)


def main():
    parser = argparse.ArgumentParser(description='Generate professional bio')
    parser.add_argument('--input', required=True, help='JSON file with bio data')
    parser.add_argument('--output', required=True, help='Output DOCX file path')
    parser.add_argument('--length', choices=['short', 'medium', 'long'],
                        default='medium', help='Bio length')
    args = parser.parse_args()

    with open(args.input, 'r') as f:
        data = json.load(f)

    builders = {
        'short': build_short_bio,
        'medium': build_medium_bio,
        'long': build_long_bio,
    }

    bio_text = builders[args.length](data)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.paragraph_format.line_spacing = 1.15

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"Professional Bio — {data.get('name', 'Unknown')}")
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(f"({args.length.capitalize()} Version)")
    sub_run.font.size = Pt(9.5)
    sub_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_paragraph('')

    for para in bio_text.split('\n\n'):
        doc.add_paragraph(para)

    doc.save(args.output)

    word_count = len(bio_text.split())
    print(json.dumps({
        "status": "success",
        "output": args.output,
        "length": args.length,
        "word_count": word_count
    }))


if __name__ == '__main__':
    main()
