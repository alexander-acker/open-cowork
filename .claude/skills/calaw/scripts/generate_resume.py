#!/usr/bin/env python3
"""Generate a professional ATS-compliant resume in DOCX format.

Usage:
    python generate_resume.py --input resume_data.json --output resume.docx [--format chronological|functional|combination]

Input JSON schema:
{
    "name": "John Doe",
    "email": "john@example.com",
    "phone": "+1 (555) 123-4567",
    "location": "San Francisco, CA",
    "linkedin": "linkedin.com/in/johndoe",
    "website": "johndoe.com",
    "summary": "Senior software engineer with 8+ years...",
    "experience": [
        {
            "title": "Senior Software Engineer",
            "company": "TechCorp",
            "location": "San Francisco, CA",
            "start_date": "Jan 2021",
            "end_date": "Present",
            "bullets": [
                "Led migration of 3 legacy systems to cloud, reducing costs by 42%",
                "Mentored team of 5 junior engineers, improving sprint velocity by 25%"
            ]
        }
    ],
    "education": [
        {
            "degree": "B.S. Computer Science",
            "school": "University of California, Berkeley",
            "graduation_date": "May 2016",
            "gpa": "3.8",
            "honors": "Magna Cum Laude"
        }
    ],
    "skills": {
        "Programming": ["Python", "JavaScript", "TypeScript", "Go"],
        "Frameworks": ["React", "Node.js", "Django"],
        "Tools": ["AWS", "Docker", "Kubernetes", "Git"]
    },
    "certifications": [
        {"name": "AWS Solutions Architect", "date": "2023"}
    ],
    "projects": [
        {
            "name": "Open Source CLI Tool",
            "description": "Built a CLI tool with 2K+ GitHub stars",
            "url": "github.com/johndoe/tool"
        }
    ]
}
"""

import argparse
import json
import sys

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Error: python-docx required. Install with: pip install python-docx", file=sys.stderr)
    sys.exit(1)


def create_styles(doc):
    """Set up document styles for ATS compliance."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    style.paragraph_format.space_after = Pt(2)
    style.paragraph_format.space_before = Pt(0)

    # Name style
    name_style = doc.styles.add_style('ResumeName', WD_STYLE_TYPE.PARAGRAPH)
    name_style.font.name = 'Calibri'
    name_style.font.size = Pt(22)
    name_style.font.bold = True
    name_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    name_style.paragraph_format.space_after = Pt(2)
    name_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Contact style
    contact_style = doc.styles.add_style('ResumeContact', WD_STYLE_TYPE.PARAGRAPH)
    contact_style.font.name = 'Calibri'
    contact_style.font.size = Pt(9.5)
    contact_style.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    contact_style.paragraph_format.space_after = Pt(6)
    contact_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Section heading style
    section_style = doc.styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
    section_style.font.name = 'Calibri'
    section_style.font.size = Pt(12)
    section_style.font.bold = True
    section_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    section_style.paragraph_format.space_before = Pt(10)
    section_style.paragraph_format.space_after = Pt(3)
    section_style.paragraph_format.border_bottom = True

    # Job title style
    job_style = doc.styles.add_style('JobTitle', WD_STYLE_TYPE.PARAGRAPH)
    job_style.font.name = 'Calibri'
    job_style.font.size = Pt(10.5)
    job_style.font.bold = True
    job_style.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    job_style.paragraph_format.space_before = Pt(6)
    job_style.paragraph_format.space_after = Pt(1)

    # Bullet style
    bullet_style = doc.styles.add_style('ResumeBullet', WD_STYLE_TYPE.PARAGRAPH)
    bullet_style.font.name = 'Calibri'
    bullet_style.font.size = Pt(10)
    bullet_style.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    bullet_style.paragraph_format.space_after = Pt(1)
    bullet_style.paragraph_format.left_indent = Inches(0.25)

    return doc


def add_section_divider(doc, title):
    """Add a section heading with underline."""
    p = doc.add_paragraph(title.upper(), style='SectionHeading')
    # Add bottom border via run formatting
    for run in p.runs:
        run.font.all_caps = True


def add_header(doc, data):
    """Add name and contact info header."""
    doc.add_paragraph(data['name'], style='ResumeName')

    contact_parts = []
    if data.get('email'):
        contact_parts.append(data['email'])
    if data.get('phone'):
        contact_parts.append(data['phone'])
    if data.get('location'):
        contact_parts.append(data['location'])
    if data.get('linkedin'):
        contact_parts.append(data['linkedin'])
    if data.get('website'):
        contact_parts.append(data['website'])

    doc.add_paragraph(' | '.join(contact_parts), style='ResumeContact')


def add_summary(doc, summary):
    """Add professional summary section."""
    if not summary:
        return
    add_section_divider(doc, 'Professional Summary')
    doc.add_paragraph(summary, style='Normal')


def add_experience(doc, experience):
    """Add work experience section."""
    if not experience:
        return
    add_section_divider(doc, 'Experience')

    for job in experience:
        # Job title and company on one line
        p = doc.add_paragraph(style='JobTitle')
        title_run = p.add_run(f"{job['title']}")
        title_run.bold = True
        p.add_run(f" | {job['company']}")

        # Date and location
        p2 = doc.add_paragraph(style='Normal')
        date_text = f"{job.get('start_date', '')} - {job.get('end_date', '')}"
        if job.get('location'):
            date_text += f" | {job['location']}"
        date_run = p2.add_run(date_text)
        date_run.font.size = Pt(9.5)
        date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        date_run.italic = True

        # Bullet points
        for bullet in job.get('bullets', []):
            bp = doc.add_paragraph(style='ResumeBullet')
            bp.add_run(f"\u2022 {bullet}")


def add_education(doc, education):
    """Add education section."""
    if not education:
        return
    add_section_divider(doc, 'Education')

    for edu in education:
        p = doc.add_paragraph(style='JobTitle')
        p.add_run(edu['degree']).bold = True
        p.add_run(f" | {edu['school']}")

        details = []
        if edu.get('graduation_date'):
            details.append(edu['graduation_date'])
        if edu.get('gpa'):
            details.append(f"GPA: {edu['gpa']}")
        if edu.get('honors'):
            details.append(edu['honors'])

        if details:
            d = doc.add_paragraph(style='Normal')
            run = d.add_run(' | '.join(details))
            run.font.size = Pt(9.5)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_skills(doc, skills):
    """Add skills section."""
    if not skills:
        return
    add_section_divider(doc, 'Skills')

    for category, skill_list in skills.items():
        p = doc.add_paragraph(style='Normal')
        cat_run = p.add_run(f"{category}: ")
        cat_run.bold = True
        p.add_run(', '.join(skill_list))


def add_certifications(doc, certs):
    """Add certifications section."""
    if not certs:
        return
    add_section_divider(doc, 'Certifications')

    for cert in certs:
        p = doc.add_paragraph(style='Normal')
        p.add_run(f"\u2022 {cert['name']}")
        if cert.get('date'):
            date_run = p.add_run(f" ({cert['date']})")
            date_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def add_projects(doc, projects):
    """Add projects section."""
    if not projects:
        return
    add_section_divider(doc, 'Projects')

    for proj in projects:
        p = doc.add_paragraph(style='JobTitle')
        p.add_run(proj['name']).bold = True
        if proj.get('url'):
            p.add_run(f" ({proj['url']})")

        if proj.get('description'):
            doc.add_paragraph(f"\u2022 {proj['description']}", style='ResumeBullet')


def generate_chronological(doc, data):
    """Standard chronological resume (most common)."""
    add_header(doc, data)
    add_summary(doc, data.get('summary'))
    add_experience(doc, data.get('experience'))
    add_education(doc, data.get('education'))
    add_skills(doc, data.get('skills'))
    add_certifications(doc, data.get('certifications'))
    add_projects(doc, data.get('projects'))


def generate_functional(doc, data):
    """Functional resume (skill-focused, for career changers)."""
    add_header(doc, data)
    add_summary(doc, data.get('summary'))
    add_skills(doc, data.get('skills'))

    # Group experience by skill area if possible
    if data.get('experience'):
        add_section_divider(doc, 'Relevant Experience')
        for job in data['experience']:
            p = doc.add_paragraph(style='Normal')
            p.add_run(f"{job['title']}, {job['company']}").bold = True
            date_text = f" ({job.get('start_date', '')} - {job.get('end_date', '')})"
            run = p.add_run(date_text)
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

            for bullet in job.get('bullets', []):
                doc.add_paragraph(f"\u2022 {bullet}", style='ResumeBullet')

    add_education(doc, data.get('education'))
    add_certifications(doc, data.get('certifications'))
    add_projects(doc, data.get('projects'))


def generate_combination(doc, data):
    """Combination resume (skills + chronological experience)."""
    add_header(doc, data)
    add_summary(doc, data.get('summary'))
    add_skills(doc, data.get('skills'))
    add_experience(doc, data.get('experience'))
    add_education(doc, data.get('education'))
    add_certifications(doc, data.get('certifications'))
    add_projects(doc, data.get('projects'))


def main():
    parser = argparse.ArgumentParser(description='Generate ATS-compliant resume')
    parser.add_argument('--input', required=True, help='JSON file with resume data')
    parser.add_argument('--output', required=True, help='Output DOCX file path')
    parser.add_argument('--format', choices=['chronological', 'functional', 'combination'],
                        default='chronological', help='Resume format')
    args = parser.parse_args()

    with open(args.input, 'r') as f:
        data = json.load(f)

    doc = Document()
    create_styles(doc)

    # Set narrow margins for more content space
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.65)
        section.right_margin = Inches(0.65)

    generators = {
        'chronological': generate_chronological,
        'functional': generate_functional,
        'combination': generate_combination,
    }

    generators[args.format](doc, data)

    doc.save(args.output)
    print(json.dumps({
        "status": "success",
        "output": args.output,
        "format": args.format,
        "sections": len(doc.paragraphs)
    }))


if __name__ == '__main__':
    main()
